from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "寒霜纪元互动影游网站产品开发文档.md"
DOCX_PATH = ROOT / "寒霜纪元互动影游网站产品开发文档.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in (("top", "80"), ("bottom", "80"), ("start", "120"), ("end", "120")):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), value)
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_cell_margins(table)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_base_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color in (
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if name != "Heading 1" else 16)
        style.paragraph_format.space_after = Pt(6)


def add_inline_runs(paragraph, text):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor.from_string("1F4D78")
        else:
            paragraph.add_run(part)


def add_code_block(doc, lines):
    text = "\n".join(lines)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    p_pr.append(shd)


def add_markdown_table(doc, lines):
    rows = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        if all(set(cell) <= {"-", ":"} for cell in cells):
            continue
        rows.append(cells)
    if not rows:
        return

    col_count = len(rows[0])
    table = doc.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"
    for row_index, row_data in enumerate(rows):
        row = table.add_row()
        for cell_index, value in enumerate(row_data):
            cell = row.cells[cell_index]
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline_runs(p, value)
            if row_index == 0:
                set_cell_shading(cell, "E8EEF5")
                for run in p.runs:
                    run.bold = True
    if col_count == 4:
        set_table_width(table, [1.15, 1.25, 1.7, 2.4])
    elif col_count == 3:
        set_table_width(table, [1.6, 1.9, 3.0])
    else:
        set_table_width(table, [6.5 / col_count] * col_count)
    doc.add_paragraph()


def build():
    doc = Document()
    set_base_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("《寒霜纪元：重生领主》互动影游网站产品开发文档")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph("版本 v1.0 / 供复刻开发使用")
    subtitle.runs[0].font.color.rgb = RGBColor.from_string("555555")
    subtitle.paragraph_format.space_after = Pt(12)

    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines = []
    table_lines = []

    for raw in lines[1:]:
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.strip().startswith("|"):
            table_lines.append(line)
            continue
        if table_lines:
            add_markdown_table(doc, table_lines)
            table_lines = []

        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=3)
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:])
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, re.sub(r"^\d+\.\s", "", stripped))
        else:
            p = doc.add_paragraph()
            add_inline_runs(p, stripped)

    if table_lines:
        add_markdown_table(doc, table_lines)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
